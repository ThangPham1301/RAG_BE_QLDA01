import logging
import csv
import json
from io import BytesIO
from io import StringIO
from django.http import StreamingHttpResponse
from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from django.views.decorators.csrf import csrf_exempt
from django.utils.decorators import method_decorator
from rest_framework import viewsets, status, permissions
from rest_framework.exceptions import PermissionDenied
from rest_framework.decorators import action
from rest_framework.views import APIView
from rest_framework.response import Response
from django.db.models import Avg, Count, Q
from django.utils import timezone
from .models import ChatSession, ChatMessage, ChatFeedback, ConversationEvaluation
from .serializers import (
    ChatSessionSerializer, ChatSessionDetailSerializer,
    ChatMessageSerializer, ChatMessageCreateSerializer, ChatFeedbackSerializer,
    ConversationEvaluationSerializer
)
from apps.realtime.events import send_notification, send_to_admins

logger = logging.getLogger(__name__)


def user_is_admin(user):
    return bool(user and user.is_authenticated and (user.is_staff or user.is_superuser))


def _build_simple_pdf(text):
    """Build a dependency-free text PDF for chat export."""
    def escape(value):
        return value.replace('\\', '\\\\').replace('(', '\\(').replace(')', '\\)')

    lines = []
    for raw_line in text.splitlines():
        line = raw_line[:110]
        while line:
            lines.append(line[:95])
            line = line[95:]
        if raw_line == '':
            lines.append('')
    if not lines:
        lines = ['No chat messages.']

    pages = []
    for start in range(0, len(lines), 42):
        page_lines = lines[start:start + 42]
        content = ['BT', '/F1 10 Tf', '50 790 Td', '14 TL']
        for index, line in enumerate(page_lines):
            if index > 0:
                content.append('T*')
            content.append(f'({escape(line)}) Tj')
        content.append('ET')
        pages.append('\n'.join(content).encode('latin-1', errors='replace'))

    objects = [b'<< /Type /Catalog /Pages 2 0 R >>']
    page_kids = []
    for idx, stream in enumerate(pages):
        page_obj_num = 3 + idx * 2
        content_obj_num = page_obj_num + 1
        page_kids.append(f'{page_obj_num} 0 R')
        objects.append(
            f'<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 842] '
            f'/Resources << /Font << /F1 << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> >> >> '
            f'/Contents {content_obj_num} 0 R >>'.encode('latin-1')
        )
        objects.append(b'<< /Length ' + str(len(stream)).encode('ascii') + b' >>\nstream\n' + stream + b'\nendstream')
    objects.insert(1, f'<< /Type /Pages /Kids [{" ".join(page_kids)}] /Count {len(pages)} >>'.encode('latin-1'))

    output = bytearray(b'%PDF-1.4\n')
    offsets = [0]
    for number, obj in enumerate(objects, start=1):
        offsets.append(len(output))
        output.extend(f'{number} 0 obj\n'.encode('ascii'))
        output.extend(obj)
        output.extend(b'\nendobj\n')
    xref_offset = len(output)
    output.extend(f'xref\n0 {len(objects) + 1}\n0000000000 65535 f \n'.encode('ascii'))
    for offset in offsets[1:]:
        output.extend(f'{offset:010d} 00000 n \n'.encode('ascii'))
    output.extend(
        f'trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF'.encode('ascii')
    )
    return bytes(output)


def _source_label(source):
    if isinstance(source, dict):
        return str(source.get('document_title') or source.get('document_id') or source)
    return str(source)


def _build_chat_docx(session, messages):
    from docx import Document

    doc = Document()
    doc.add_heading(session.title or 'Chat export', level=1)
    doc.add_paragraph(f'Project: {session.project.name}')
    doc.add_paragraph(f'Created: {session.created_at.isoformat()}')
    doc.add_paragraph(f'Updated: {session.updated_at.isoformat()}')

    for message in messages:
        doc.add_heading(f'{message.role.upper()} - {message.created_at.isoformat()}', level=2)
        doc.add_paragraph(message.content or '')
        if message.sources:
            doc.add_paragraph('Sources: ' + ', '.join(_source_label(source) for source in message.sources if source))

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def build_chat_export_response(session, request):
    export_format = request.query_params.get('format', 'docx').lower()
    if export_format == 'doc':
        export_format = 'docx'
    if export_format not in {'docx', 'pdf', 'csv', 'txt', 'json'}:
        return Response({'error': 'format must be docx, pdf, csv, txt, or json'}, status=status.HTTP_400_BAD_REQUEST)

    messages = list(session.messages.order_by('created_at'))
    filename_base = f'chat-{session.id}'

    payload = {
        'id': session.id,
        'title': session.title,
        'project': session.project.name,
        'created_at': session.created_at.isoformat(),
        'updated_at': session.updated_at.isoformat(),
        'messages': ChatMessageSerializer(messages, many=True).data,
    }

    if export_format == 'docx':
        response = HttpResponse(
            _build_chat_docx(session, messages),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        response['Content-Disposition'] = f'attachment; filename="{filename_base}.docx"'
        return response

    if export_format == 'json':
        response = HttpResponse(
            json.dumps(payload, ensure_ascii=False, indent=2),
            content_type='application/json; charset=utf-8',
        )
        response['Content-Disposition'] = f'attachment; filename="{filename_base}.json"'
        return response

    if export_format == 'csv':
        buffer = StringIO()
        writer = csv.writer(buffer)
        writer.writerow(['Chat title', session.title])
        writer.writerow(['Project', session.project.name])
        writer.writerow([])
        writer.writerow(['Created at', 'Role', 'Content', 'Sources'])
        for message in messages:
            sources = '; '.join(
                _source_label(source)
                for source in (message.sources or [])
                if source
            )
            writer.writerow([message.created_at.isoformat(), message.role, message.content, sources])
        response = HttpResponse(buffer.getvalue(), content_type='text/csv; charset=utf-8')
        response['Content-Disposition'] = f'attachment; filename="{filename_base}.csv"'
        return response

    lines = [
        f'Chat: {session.title}',
        f'Project: {session.project.name}',
        f'Created: {session.created_at.isoformat()}',
        '',
    ]
    for message in messages:
        lines.append(f'[{message.created_at.isoformat()}] {message.role.upper()}')
        lines.append(message.content)
        if message.sources:
            source_labels = [
                _source_label(source)
                for source in message.sources
                if source
            ]
            lines.append(f'Sources: {", ".join(source_labels)}')
        lines.append('')
    text = '\n'.join(lines)

    if export_format == 'txt':
        response = HttpResponse(text, content_type='text/plain; charset=utf-8')
        response['Content-Disposition'] = f'attachment; filename="{filename_base}.txt"'
        return response

    response = HttpResponse(_build_simple_pdf(text), content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{filename_base}.pdf"'
    return response


class ChatSessionViewSet(viewsets.ModelViewSet):
    """Chat session management"""
    serializer_class = ChatSessionSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        queryset = ChatSession.objects.filter(
            user=self.request.user,
            is_deleted=False
        ).order_by('-updated_at')

        project_id = self.request.query_params.get('project_id')
        if project_id:
            queryset = queryset.filter(project_id=project_id)

        return queryset

    def get_serializer_class(self):
        if self.action == 'retrieve':
            return ChatSessionDetailSerializer
        return ChatSessionSerializer

    def create(self, request, *args, **kwargs):
        """Create new chat session"""
        data = request.data.copy()
        project_id = data.get('project_id')
        if not project_id:
            return Response({'error': 'project_id is required'}, status=status.HTTP_400_BAD_REQUEST)
        from apps.projects.models import Project
        if not Project.objects.filter(id=project_id, owner=request.user).exists():
            return Response({'error': 'project is not available'}, status=status.HTTP_400_BAD_REQUEST)
        data['project'] = project_id
        data['user'] = request.user.id
        serializer = self.get_serializer(data=data)
        serializer.is_valid(raise_exception=True)
        self.perform_create(serializer)
        return Response(serializer.data, status=status.HTTP_201_CREATED)

    def perform_create(self, serializer):
        serializer.save(user=self.request.user)

    def perform_update(self, serializer):
        serializer.save()

    def destroy(self, request, *args, **kwargs):
        """Delete session and cascade its documents"""
        instance = self.get_object()
        instance.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)

    @action(detail=True, methods=['post'])
    def send_message(self, request, pk=None):
        """Send message to chat session and get RAG response"""
        session = self.get_object()
        
        serializer = ChatMessageCreateSerializer(data=request.data)
        if not serializer.is_valid():
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

        content = serializer.validated_data['content']

        try:
            # Import here to avoid circular imports
            from .chat_service import ChatService

            chat_service = ChatService()

            # document_id optional — cho phép hỏi về 1 file cụ thể
            document_id = request.data.get('document_id')
            document_id = int(document_id) if document_id else None
            if document_id:
                from apps.documents.models import Document
                from apps.teams.permissions import user_can_access_document
                document = Document.objects.filter(id=document_id, is_deleted=False).first()
                if not document or not user_can_access_document(request.user, document):
                    return Response({'error': 'document is not available'}, status=status.HTTP_403_FORBIDDEN)

            # Get RAG response
            response_data = chat_service.ask_question(session.id, content, document_id=document_id)

            # Update session timestamp
            session.last_message_at = timezone.now()
            session.save(update_fields=['last_message_at'])

            return Response(response_data, status=status.HTTP_200_OK)

        except Exception as e:
            logger.error(f"Error in send_message: {e}", exc_info=True)
            return Response(
                {'error': str(e)},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

    @action(detail=True, methods=['get'])
    def messages(self, request, pk=None):
        """Get all messages in session"""
        session = self.get_object()
        messages = session.messages.order_by('created_at')
        serializer = ChatMessageSerializer(messages, many=True)
        return Response(serializer.data)

    @action(detail=True, methods=['get'], url_path='export')
    def export(self, request, pk=None):
        """Export the full chat session including user questions, assistant answers, and sources."""
        session = self.get_object()
        return build_chat_export_response(session, request)


class ChatSessionExportView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request, pk=None):
        session = get_object_or_404(ChatSession, id=pk, user=request.user, is_deleted=False)
        return build_chat_export_response(session, request)


class ChatSendView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request):
        session_id = request.data.get('chat_session_id')
        if not session_id:
            return Response({'error': 'chat_session_id is required'}, status=status.HTTP_400_BAD_REQUEST)

        session = get_object_or_404(ChatSession, id=session_id, user=request.user, is_deleted=False)
        serializer = ChatMessageCreateSerializer(data=request.data)
        serializer.is_valid(raise_exception=True)

        try:
            from .chat_service import ChatService

            chat_service = ChatService()

            # document_id optional — cho phép hỏi về 1 file cụ thể
            document_id = request.data.get('document_id')
            document_id = int(document_id) if document_id else None
            if document_id:
                from apps.documents.models import Document
                from apps.teams.permissions import user_can_access_document
                document = Document.objects.filter(id=document_id, is_deleted=False).first()
                if not document or not user_can_access_document(request.user, document):
                    return Response({'error': 'document is not available'}, status=status.HTTP_403_FORBIDDEN)

            response_data = chat_service.ask_question(
                session.id,
                serializer.validated_data['content'],
                document_id=document_id,
            )
            session.last_message_at = timezone.now()
            session.save(update_fields=['last_message_at'])
            
            # Serialize message object properly
            msg = response_data.get('message')
            if msg:
                # Get fresh message instance to ensure proper serialization
                from .models import ChatMessage
                msg_obj = ChatMessage.objects.get(id=msg['id'])
                msg_serializer = ChatMessageSerializer(msg_obj)
                response_data['message'] = msg_serializer.data
            
            return Response(response_data, status=status.HTTP_200_OK)
        except Exception as e:
            logger.error(f'Error in chat send: {e}', exc_info=True)
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class ChatStreamView(APIView):
    """SSE streaming endpoint — /api/chat/stream/

    Dùng Server-Sent Events (SSE) để push từng token LLM về FE ngay khi sinh ra.
    FE dùng fetch() + ReadableStream để đọc, không dùng EventSource
    (vì EventSource không support POST và auth header).
    """
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request):
        session_id = request.data.get('chat_session_id')
        content = request.data.get('content', '').strip()

        if not session_id:
            return Response({'error': 'chat_session_id is required'}, status=status.HTTP_400_BAD_REQUEST)
        if not content:
            return Response({'error': 'content is required'}, status=status.HTTP_400_BAD_REQUEST)

        session = get_object_or_404(ChatSession, id=session_id, user=request.user, is_deleted=False)

        document_id = request.data.get('document_id')
        document_id = int(document_id) if document_id else None
        if document_id:
            from apps.documents.models import Document
            from apps.teams.permissions import user_can_access_document
            document = Document.objects.filter(id=document_id, is_deleted=False).first()
            if not document or not user_can_access_document(request.user, document):
                return Response({'error': 'document is not available'}, status=status.HTTP_403_FORBIDDEN)

        from .chat_service import ChatService
        chat_service = ChatService()

        def event_stream():
            yield 'retry: 3000\n\n'  # FE retry sau 3s nếu connection drop
            try:
                for sse_event in chat_service.ask_question_stream(
                    session_id=session.id,
                    question=content,
                    document_id=document_id,
                ):
                    yield sse_event
            except Exception as exc:
                import json
                logger.error('ChatStreamView: unhandled error: %s', exc, exc_info=True)
                yield f'data: {json.dumps({"type": "error", "content": str(exc)[:80]})}\n\n'

        response = StreamingHttpResponse(
            event_stream(),
            content_type='text/event-stream',
        )
        response['Cache-Control'] = 'no-cache'
        response['X-Accel-Buffering'] = 'no'  # tắt buffering ở nginx
        return response


class ChatMessageViewSet(viewsets.ReadOnlyModelViewSet):
    """Read-only chat messages"""
    serializer_class = ChatMessageSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        session_id = self.request.query_params.get('session_id')
        if session_id:
            return ChatMessage.objects.filter(
                chat_session_id=session_id,
                chat_session__user=self.request.user,
                chat_session__is_deleted=False,
            ).order_by('created_at')
        return ChatMessage.objects.none()


class ChatFeedbackViewSet(viewsets.ModelViewSet):
    """Chat feedback management"""
    serializer_class = ChatFeedbackSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        return ChatFeedback.objects.filter(user=self.request.user)

    def perform_create(self, serializer):
        serializer.save(user=self.request.user)


class ConversationEvaluationViewSet(viewsets.ModelViewSet):
    """Conversation-level evaluation. One official evaluation per chat session."""
    serializer_class = ConversationEvaluationSerializer
    permission_classes = [permissions.IsAuthenticated]

    def _send_admin_evaluation_event(self, event_type, data):
        send_to_admins(event_type, data)

        from django.contrib.auth import get_user_model
        from apps.teams.models import InAppNotification

        User = get_user_model()
        title = 'New chat evaluation' if event_type == 'evaluation.created' else 'Chat evaluation updated'
        chat_label = data.get('chat_title') or f"session #{data.get('chat_session')}"
        message = (
            f"{data.get('user_email') or 'A user'} rated "
            f"{chat_label} "
            f"{data.get('rating')}/5."
        )
        for admin in User.objects.filter(is_active=True).filter(Q(is_staff=True) | Q(is_superuser=True)):
            notification = InAppNotification.objects.create(
                user=admin,
                title=title,
                message=message,
                data={
                    'type': event_type,
                    'evaluation_id': data.get('id'),
                    'chat_session_id': data.get('chat_session'),
                    'rating': data.get('rating'),
                },
            )
            send_notification(notification)

    def get_queryset(self):
        queryset = ConversationEvaluation.objects.select_related(
            'chat_session',
            'chat_session__project',
            'user',
            'pinned_by',
        )
        if not user_is_admin(self.request.user):
            queryset = queryset.filter(user=self.request.user)

        chat_session_id = self.request.query_params.get('chat_session')
        if chat_session_id:
            queryset = queryset.filter(chat_session_id=chat_session_id)

        rating = self.request.query_params.get('rating')
        if rating:
            queryset = queryset.filter(rating=rating)

        pinned = self.request.query_params.get('pinned')
        if pinned in {'true', 'false'}:
            queryset = queryset.filter(is_pinned=(pinned == 'true'))

        return queryset.order_by('-is_pinned', '-updated_at')

    def create(self, request, *args, **kwargs):
        chat_session_id = request.data.get('chat_session')
        if not chat_session_id:
            return Response({'chat_session': 'This field is required.'}, status=status.HTTP_400_BAD_REQUEST)

        chat_session = ChatSession.objects.filter(
            id=chat_session_id,
            user=request.user,
            is_deleted=False,
        ).first()
        if not chat_session:
            return Response({'chat_session': 'Chat session is not available.'}, status=status.HTTP_400_BAD_REQUEST)

        existing = ConversationEvaluation.objects.filter(chat_session=chat_session, user=request.user).first()
        if existing:
            update_serializer = self.get_serializer(existing, data=request.data, partial=True)
            update_serializer.is_valid(raise_exception=True)
            self._save_user_update(update_serializer)
            self._send_admin_evaluation_event('evaluation.updated', update_serializer.data)
            return Response(update_serializer.data, status=status.HTTP_200_OK)

        serializer = self.get_serializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        serializer.save(user=request.user)
        self._send_admin_evaluation_event('evaluation.created', serializer.data)
        return Response(serializer.data, status=status.HTTP_201_CREATED)

    def perform_update(self, serializer):
        if serializer.instance.user_id != self.request.user.id:
            raise PermissionDenied('Only the owner can edit this evaluation.')
        self._save_user_update(serializer)
        self._send_admin_evaluation_event('evaluation.updated', serializer.data)

    def destroy(self, request, *args, **kwargs):
        if user_is_admin(request.user):
            return Response({'detail': 'Admin cannot delete user evaluations.'}, status=status.HTTP_403_FORBIDDEN)
        return super().destroy(request, *args, **kwargs)

    def _save_user_update(self, serializer):
        instance = serializer.instance
        if instance and instance.is_pinned:
            serializer.save(
                user=self.request.user,
                is_pinned=False,
                pinned_at=None,
                pinned_by=None,
            )
        else:
            serializer.save(user=self.request.user)

    @action(detail=True, methods=['post'])
    def pin(self, request, pk=None):
        if not user_is_admin(request.user):
            return Response({'detail': 'Admin role is required.'}, status=status.HTTP_403_FORBIDDEN)
        evaluation = self.get_object()
        evaluation.is_pinned = True
        evaluation.pinned_at = timezone.now()
        evaluation.pinned_by = request.user
        evaluation.save(update_fields=['is_pinned', 'pinned_at', 'pinned_by', 'updated_at'])
        data = self.get_serializer(evaluation).data
        send_to_admins('evaluation.pinned', data)
        return Response(data)

    @action(detail=True, methods=['post'])
    def unpin(self, request, pk=None):
        if not user_is_admin(request.user):
            return Response({'detail': 'Admin role is required.'}, status=status.HTTP_403_FORBIDDEN)
        evaluation = self.get_object()
        evaluation.is_pinned = False
        evaluation.pinned_at = None
        evaluation.pinned_by = None
        evaluation.save(update_fields=['is_pinned', 'pinned_at', 'pinned_by', 'updated_at'])
        data = self.get_serializer(evaluation).data
        send_to_admins('evaluation.unpinned', data)
        return Response(data)

    @action(detail=False, methods=['get'])
    def stats(self, request):
        if not user_is_admin(request.user):
            return Response({'detail': 'Admin role is required.'}, status=status.HTTP_403_FORBIDDEN)

        queryset = self.get_queryset()
        distribution = {
            row['rating']: row['total']
            for row in queryset.values('rating').annotate(total=Count('id')).order_by('rating')
        }
        return Response({
            'total': queryset.count(),
            'pinned': queryset.filter(is_pinned=True).count(),
            'average_rating': round(queryset.aggregate(value=Avg('rating'))['value'] or 0, 2),
            'distribution': {str(score): distribution.get(score, 0) for score in range(1, 6)},
        })
