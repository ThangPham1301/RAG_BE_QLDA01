"""Document parsing utilities (PDF extraction)."""
from __future__ import annotations

import logging
import re
import unicodedata
from typing import List
from decouple import config
try:
    from PIL import Image
    import pytesseract
    # Configure path for Windows users via .env, fallback to standard Tesseract path
    pytesseract.pytesseract.tesseract_cmd = config(
        'TESSERACT_CMD', 
        default=r'C:\Program Files\Tesseract-OCR\tesseract.exe'
    )
except ImportError:
    pass

logger = logging.getLogger(__name__)


def _get_fitz():
    """Import PyMuPDF lazily so a running server can pick up newly installed packages."""
    try:
        import fitz  # pymupdf
        return fitz
    except Exception as exc:  # pragma: no cover
        logger.warning(f'PyMuPDF (fitz) import failed: {exc}')
        return None


def extract_text_from_pdf(path: str) -> str:
    """Extract text from a PDF using PyMuPDF. Returns plain text.

    If PyMuPDF is not available, returns empty string.
    """
    logger.debug(f'[extract_text_from_pdf] Starting for {path}')
    fitz = _get_fitz()
    if fitz is None:
        logger.error('[extract_text_from_pdf] PyMuPDF not available')
        return ''
    
    try:
        doc = fitz.open(path)
        logger.debug(f'[extract_text_from_pdf] PDF opened, pages={len(doc)}')
    except Exception as e:
        logger.error(f'[extract_text_from_pdf] Failed to open PDF: {e}', exc_info=True)
        return ''
    
    parts: List[str] = []
    has_text = False
    
    for idx, page in enumerate(doc):
        try:
            # 1. Thử lấy text thông thường trước
            text = (page.get_text("text") or page.get_text()).strip()
            
            # 2. Nếu trang trống hoặc cực ít chữ -> Khả năng cao là ảnh scan -> dùng OCR
            # Tăng ngưỡng lên 500 ký tự vì chữ ký số điện tử (red seal) thường chứa
            # 1 layer text ẩn khoảng 100-200 ký tự. Nếu để 50, PyMuPDF sẽ tưởng
            # đã trích xuất thành công và bỏ qua toàn bộ phần ảnh của văn bản.
            if len(text) < 500:
                logger.info(f'[extract_text_from_pdf] Page {idx} has very little text ({len(text)} chars). Falling back to OCR...')
                try:
                    # Chuyển trang PDF thành hình ảnh
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # Zoom 2x để rõ chữ hơn
                    if pix.alpha:
                        pix = fitz.Pixmap(fitz.csRGB, pix)
                    
                    # Convert sang PIL Image
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    
                    # Chạy Tesseract OCR (mặc định lấy cả tiếng Việt và Anh)
                    ocr_text = pytesseract.image_to_string(img, lang="vie+eng").strip()
                    if ocr_text:
                        text += f"\n{ocr_text}"
                except Exception as ocr_exc:
                    logger.warning(f'[extract_text_from_pdf] OCR failed on page {idx}: {ocr_exc}')

            if text.strip():
                parts.append(text)
                logger.debug(f'[extract_text_from_pdf] Page {idx}: {len(text)} chars')
                has_text = True
                
        except Exception as e:
            logger.warning(f'[extract_text_from_pdf] Failed to extract page {idx}: {e}')
    
    result = '\n\n'.join(parts)
    logger.debug(f'[extract_text_from_pdf] Total extracted: {len(result)} chars')
    return result

def extract_text_from_image(path: str) -> str:
    """Extract text from an image file (png, jpg, jpeg) using Tesseract OCR."""
    try:
        img = Image.open(path)
        logger.info(f'[extract_text_from_image] Running OCR on {path}')
        text = pytesseract.image_to_string(img, lang="vie+eng").strip()
        logger.debug(f'[extract_text_from_image] Total extracted: {len(text)} chars')
        return text
    except Exception as exc:
        logger.error(f'[extract_text_from_image] OCR Extraction failed: {exc}', exc_info=True)
        return ''


def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
    """Naive chunking by characters.

    Returns list of text chunks with overlap.
    """
    if not text:
        return []
    chunks: List[str] = []
    start = 0
    L = len(text)
    while start < L:
        end = min(start + chunk_size, L)
        chunks.append(text[start:end])
        # Ensure we move forward; if end == L, we're done
        if end == L:
            break
        start = max(start + 1, end - overlap)  # Ensure forward progress
    return chunks


def _fold_text(text: str) -> str:
    text = unicodedata.normalize("NFD", text or "").lower()
    text = "".join(ch for ch in text if unicodedata.category(ch) != "Mn")
    return text.replace("đ", "d")


def _is_numbered_item(line: str) -> bool:
    return bool(re.match(r"^\s*(?:\d+|[ivxlcdm]+)\s*[\.\)]\s+", _fold_text(line)))


def _is_signature_title(line: str) -> bool:
    folded = _fold_text(line)
    return any(
        marker in folded
        for marker in [
            "kt.",
            "tm.",
            "tl.",
            "tuq.",
            "pho chu nhiem",
            "chu nhiem",
            "bo truong",
            "thu truong",
            "chu tich",
            "giam doc",
        ]
    )


def _line_section_type(line: str) -> str | None:
    folded = _fold_text(line).strip()
    if not folded:
        return None
    if "kinh gui" in folded:
        return "kinh_gui"
    if folded.startswith("noi nhan") or folded.startswith("noi nhận"):
        return "noi_nhan"
    if _is_signature_title(line):
        return "chu_ky"
    if _is_numbered_item(line):
        return "muc_danh_so"
    if folded.startswith("can cu") or folded.startswith("xet de nghi"):
        return "can_cu"
    if folded.startswith("van phong") and "thong bao" in folded:
        return "ket_luan"
    return None


def _clean_lines(text: str) -> List[str]:
    lines = [re.sub(r"\s+", " ", line).strip() for line in (text or "").splitlines()]
    return [line for line in lines if line]


def _header_context(lines: List[str], max_chars: int = 700) -> str:
    keep = []
    for line in lines[:25]:
        section_type = _line_section_type(line)
        if section_type in {"kinh_gui", "can_cu", "muc_danh_so"}:
            break
        keep.append(line)
        if sum(len(item) + 1 for item in keep) >= max_chars:
            break
    return "\n".join(keep).strip()


def _pack_paragraphs(
    paragraphs: List[str],
    chunk_type: str,
    chunk_size: int,
    context: str = "",
    context_score: str = "high",
) -> List[dict]:
    chunks: List[dict] = []
    current: List[str] = []

    def flush() -> None:
        if not current:
            return
        body = "\n".join(current).strip()
        prefix = f"{context}\n\n" if context and context not in body else ""
        chunks.append(
            {
                "chunk": f"{prefix}{body}".strip(),
                "chunk_type": chunk_type,
                "context_score": context_score,
            }
        )
        current.clear()

    for paragraph in paragraphs:
        paragraph = paragraph.strip()
        if not paragraph:
            continue
        candidate = "\n".join(current + [paragraph])
        extra_context = len(context) + 2 if context and context not in candidate else 0
        if current and len(candidate) + extra_context > chunk_size:
            flush()
        current.append(paragraph)

    flush()
    return chunks


def chunk_administrative_text(text: str, chunk_size: int = 1200) -> List[dict]:
    """Rule-based chunking for Vietnamese administrative documents.

    Keeps legal/admin anchors together without using an LLM. Falls back to
    character chunks if the text does not look like an administrative document.
    """
    lines = _clean_lines(text)
    if not lines:
        return []

    admin_hits = sum(
        1
        for line in lines[:80]
        if _line_section_type(line)
        or any(
            marker in _fold_text(line)
            for marker in [
                "cong hoa xa hoi",
                "doc lap",
                "so:",
                "ngay",
                "van phong",
                "bo ",
                "uy ban",
            ]
        )
    )
    if admin_hits < 2:
        return [
            {"chunk": chunk, "chunk_type": "Khac", "context_score": "medium"}
            for chunk in chunk_text(text, chunk_size=chunk_size, overlap=200)
        ]

    context = _header_context(lines)
    chunks: List[dict] = []
    header_lines: List[str] = []
    current_type = "noi_dung"
    current_lines: List[str] = []

    def flush_current() -> None:
        nonlocal current_lines
        if not current_lines:
            return
        chunks.extend(_pack_paragraphs(current_lines, current_type, chunk_size, context=context))
        current_lines = []

    in_body = False
    for line in lines:
        section_type = _line_section_type(line)

        if not in_body:
            if section_type in {"kinh_gui", "can_cu", "muc_danh_so"}:
                in_body = True
                if header_lines:
                    chunks.extend(_pack_paragraphs(header_lines, "header", chunk_size, context_score="high"))
                header_lines = []
                current_type = section_type
                current_lines = [line]
                continue
            header_lines.append(line)
            continue

        if section_type in {"kinh_gui", "can_cu", "noi_nhan", "chu_ky", "ket_luan"}:
            if current_type == section_type and section_type in {"noi_nhan", "chu_ky"}:
                current_lines.append(line)
                continue
            flush_current()
            current_type = section_type
            current_lines = [line]
            continue

        if section_type == "muc_danh_so":
            flush_current()
            current_type = "muc_danh_so"
            current_lines = [line]
            continue

        current_lines.append(line)

    if header_lines and not chunks:
        chunks.extend(_pack_paragraphs(header_lines, "header", chunk_size, context_score="medium"))
    flush_current()

    # Very small documents can produce only one huge chunk; keep a deterministic fallback.
    normalized_chunks = [item for item in chunks if item.get("chunk", "").strip()]
    if not normalized_chunks:
        return [
            {"chunk": chunk, "chunk_type": "Khac", "context_score": "medium"}
            for chunk in chunk_text(text, chunk_size=chunk_size, overlap=200)
        ]
    return normalized_chunks


def extract_text_from_docx(path: str) -> str:
    """Extract text from a DOCX including paragraphs and tables."""
    try:
        from docx import Document as DocxDocument
    except Exception as exc:
        logger.error(f'[extract_text_from_docx] python-docx not available: {exc}')
        return ''

    try:
        doc = DocxDocument(path)
    except Exception as exc:
        logger.error(f'[extract_text_from_docx] Failed to open DOCX: {exc}', exc_info=True)
        return ''

    parts: List[str] = []

    for paragraph in doc.paragraphs:
        text = (paragraph.text or '').strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if row_text:
                parts.append(' | '.join(row_text))

    result = '\n'.join(parts)
    logger.debug(f'[extract_text_from_docx] Total extracted: {len(result)} chars')
    return result
